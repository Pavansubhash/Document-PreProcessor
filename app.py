from flask import Flask, render_template, request, send_file
from PIL import Image, ImageEnhance
import easyocr
import os
import fitz  # PyMuPDF
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from fpdf import FPDF

def convert_to_pdf(text, output_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, text)
    pdf.output(output_path)

app = Flask(__name__)

ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'gif', 'docx'}

reader = easyocr.Reader(['en'], gpu=False)  # Adjust languages and GPU usage as needed

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ''
    for page_number in range(doc.page_count):
        page = doc[page_number]
        text += page.get_text()
    return text

def enhance_image_quality(image_path):
    image = Image.open(image_path)
    # Additional preprocessing: Convert the image to grayscale and adjust contrast
    image = image.convert('L')
    enhanced_image = ImageEnhance.Contrast(image).enhance(2.0)  # Adjust contrast (you can modify as needed)
    return enhanced_image

def extract_text_from_image(image_path):
    img = Image.open(image_path)
    text_result = reader.readtext(image_path)
    
    text = ""
    for detection in text_result:
        text += detection[1] + ' '

    return text

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

def convert_to_docx(text, output_path):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_path)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process_document', methods=['POST'])
def process_document():
    if 'file' not in request.files:
        return render_template('index.html', error='No file part')

    file = request.files['file']
    destination_location = request.form.get('destination_location', '.')

    if file.filename == '':
        return render_template('index.html', error='No selected file')

    if file and allowed_file(file.filename):
        file_path = os.path.join(destination_location, file.filename)
        file.save(file_path)

        text = ''

        if file.filename.lower().endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
            text = extract_text_from_image(file_path)
        elif file.filename.lower().endswith('.docx'):
            text = extract_text_from_docx(file_path)

        # If text is not recognized well for images, enhance the image and try again
        if not text and file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
            enhanced_image = enhance_image_quality(file_path)
            enhanced_image_path = os.path.join(destination_location, 'enhanced_image.png')
            enhanced_image.save(enhanced_image_path)
            text = extract_text_from_image(enhanced_image_path)

        format_choice = request.form.get('format_choice', 'txt')

        # Save the extracted text to a file
        text_file_path = os.path.join(destination_location, 'extracted_text.txt')
        with open(text_file_path, 'w', encoding='utf-8') as text_file:
            text_file.write(text)

        # Convert to the selected format
        output_file_path = os.path.join(destination_location, f'extracted_text.{format_choice}')
        if format_choice == 'docx':
            convert_to_docx(text, output_file_path)
        elif format_choice == 'pdf':
            convert_to_pdf(text, output_file_path)

        # Provide the file for download
        return send_file(output_file_path, as_attachment=True, download_name=f'extracted_text.{format_choice}')

    return render_template('index.html', error='Invalid file format')

if __name__ == '__main__':
    app.run(debug=True)
